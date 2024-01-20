from flask import request, make_response, current_app
from server.database.db import db
from server.database.models import Document
from server.database.schema import DocumentSchema
from sqlalchemy.sql import func
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

import requests
import os
import logging

logging.basicConfig(level=logging.DEBUG, 
                    format='%(asctime)s - %(levelname)s - %(message)s')


document_schema = DocumentSchema()
documents_schema = DocumentSchema(many=True)

"""
    read_and_parse_docx function reads the content of the upload file and divide it into paragraphs.
    :param 1: path of the file in server/upload_file
    :return: divided paragraphs from the file contents
"""
def read_and_parse_docx(file_path):
    try:
        logging.debug(f"Reading DOCX file: {file_path}")
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip() != '':
                paragraph_data = {
                    "text": para.text,
                    "runs": [{"text": run.text, "style": run.style} for run in para.runs]
                }
                paragraphs.append(paragraph_data)
    except Exception as e:
        logging.error(f"Error in read_and_parse_docx: {e}")
        raise
    return paragraphs

"""
    correct_text_with_api function work with the LLM API to correct a paragraph.
    :param 1: a paragraph
    :return: corrected sentences of that paragraph
"""
def correct_text_with_api(paragraph):
    api_url = "https://polite-horribly-cub.ngrok-free.app/generate_code"
    prompt = f"Correct English of this text: {paragraph}. Here is the corrected version:"

    try:
        # Send the GET request with the prompt
        response = requests.get(api_url, params={"prompts": prompt})
        if response.status_code == 200:
            data = response.json()
            corrected_paragraph = data[0]
            corrected_sentence = corrected_paragraph.split('\n')[0]  # Extracts the first sentence
            return corrected_sentence
        else:
            print(f"Failed to correct text. Status code: {response.status_code}")
            return paragraph
    except Exception as e:
        # Handle exceptions for the API request
        print(f"An exception occurred: {e}")
        return paragraph

"""
    create_corrected_docx function will create a new document file with corrected content.
    :param 1: path of the upload file
    :param 2: corrected paragraphs
    :return: path of the corrected file
"""
# def create_corrected_docx(original_file_path, corrected_paragraphs):
#     try:
#         logging.debug(f"Creating corrected DOCX file from: {original_file_path}")
#         doc = DocxDocument(original_file_path)
#         for i, para in enumerate(doc.paragraphs):
#             if i < len(corrected_paragraphs):
#                 # Remember original paragraph formatting
#                 original_alignment = para.paragraph_format.alignment

#                 # Clear all runs in the paragraph
#                 for run in para.runs:
#                     run.clear()

#                 # Rebuild the paragraph with corrected text but original formatting
#                 corrected_para_text = corrected_paragraphs[i].split()
#                 for word in corrected_para_text:
#                     # Add run with original formatting
#                     new_run = para.add_run(word + " ")
#                     if para.runs:  # Copy style from the first run if available
#                         original_run = para.runs[0]
#                         new_run.bold = original_run.bold
#                         new_run.italic = original_run.italic
#                         new_run.underline = original_run.underline
#                         new_run.font.size = original_run.font.size
#                         new_run.font.name = original_run.font.name

#                 # Apply the original paragraph formatting
#                 para.paragraph_format.alignment = original_alignment

#         corrected_file_name = os.path.basename(original_file_path).replace('.docx', '_corrected.docx')
#         corrected_file_path = os.path.join(current_app.config['DOWNLOAD_FOLDER'], corrected_file_name)
#         doc.save(corrected_file_path)
#     except Exception as e:
#         logging.error(f"Error in create_corrected_docx: {e}")
#         raise
#     return corrected_file_path

def get_run_formatting(run):
    """Extracts the formatting of a run."""
    return {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_name': run.font.name,
        'font_size': run.font.size,
        'font_color': run.font.color.rgb if run.font.color else None,
        'highlight_color': run.font.highlight_color,
        'strike': run.font.strike,
        'subscript': run.font.subscript,
        'superscript': run.font.superscript,
        # Add more formatting properties as needed
    }

def apply_run_formatting(run, formatting):
    """Applies the formatting to a run."""
    run.bold = formatting.get('bold')
    run.italic = formatting.get('italic')
    run.underline = formatting.get('underline')
    run.font.name = formatting.get('font_name')
    run.font.size = formatting.get('font_size')
    if formatting.get('font_color') is not None:
        run.font.color.rgb = formatting.get('font_color')
    run.font.highlight_color = formatting.get('highlight_color')
    run.font.strike = formatting.get('strike')
    run.font.subscript = formatting.get('subscript')
    run.font.superscript = formatting.get('superscript')
    # Apply more formatting properties as needed

def create_corrected_docx(original_file_path, corrected_paragraphs):
    try:
        logging.debug(f"Creating corrected DOCX file from: {original_file_path}")
        doc = DocxDocument(original_file_path)

        for i, para in enumerate(doc.paragraphs):
            if i < len(corrected_paragraphs):
                # Extract original run formatting
                run_formats = [get_run_formatting(run) for run in para.runs]

                # Clear the paragraph
                para.clear()

                # Rebuild the paragraph with corrected text
                corrected_para = corrected_paragraphs[i]
                new_run = para.add_run(corrected_para)

                # Apply the original formatting to the new runs
                for fmt in run_formats:
                    apply_run_formatting(new_run, fmt)

        corrected_file_name = os.path.basename(original_file_path).replace('.docx', '_corrected.docx')
        corrected_file_path = os.path.join(current_app.config['DOWNLOAD_FOLDER'], corrected_file_name)
        doc.save(corrected_file_path)
    except Exception as e:
        logging.error(f"Error in create_corrected_docx: {e}")
        raise
    return corrected_file_path

# Get all document by user ID
def get_all_documents_of_user_service(user_id, page, per_page):
    try:
        documents = Document.query.filter_by(user_id=user_id).paginate(page=page, per_page=per_page, error_out=False)

        if documents.items:
            return make_response(
                documents_schema.jsonify(documents.items),
                200
            )
        else:
            return make_response(
                {"message": "No documents found"},
                404  # Not Found
            )
    except Exception as e:
        print(e)
        db.session.rollback()
        return make_response(
            {"message": "Unable to find all documents"},
            400  # Bad Request
        )
# def get_all_documents_of_user_service(user_id):
#     try:
#         documents = Document.query.filter_by(user_id=user_id).all()
#         if documents:
#             return make_response(
#                 documents_schema.jsonify(documents),
#                 200
#             )
#         else:
#             return make_response(
#                 {"message": "No documents found"},
#                 404 # Not Found
#             )
#
#     except Exception as e:
#         print(e)
#         db.session.rollback()
#         return make_response(
#             {"message": "Unable to find all documents"},
#             400 # Bad Request
#         )

# Get document of a user by document ID
def get_document_of_user_by_id_service(user_id, document_id):
    try:
        document = Document.query.filter_by(user_id=user_id, document_id=document_id).first()
        if document:
            return make_response(
                document_schema.jsonify(document), 
                200
            )
        else:
            return make_response(
                {"message": "Document not found"},
                404 # Not Found
            )


    except Exception as e:
        print(e)
        db.session.rollback()
        return make_response(
            {"message": "Unable to find document by id: " + document_id},
            400 # Bad Request
        )
    
# Get document of a user by document name
def get_document_of_user_by_name_service(user_id, document_name):
    try:
        data = document_name or request.args.get('name')

        if data:
            documents = Document.query.filter(
                Document.user_id == user_id,
                func.lower(Document.document_name).like('%' + data.lower() + '%')
            ).all()
            if documents:
                return make_response(
                    documents_schema.jsonify(documents),
                    200
                )
            else:
                return make_response(
                    {"message": "Document not found"},
                    404 # Not Found
                )
        else:
            return make_response(
                {"message": "Invalid request"},
                400 # Bad Request
            )
        
    except Exception as e:
        print(e)
        db.session.rollback()
        return make_response(
            {"message": "Unable to find document by name: " + str(document_name)},
            400 # Bad Request
        )